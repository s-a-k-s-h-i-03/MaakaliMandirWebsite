from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image


ROOT = Path(r"C:\Users\saksh\maakalisonkundgpm.org_new")
SCREENSHOT_DIR = Path(r"C:\Users\saksh\OneDrive\Pictures\Screenshots")
OUTPUT = ROOT / "deliverables" / "mandir-website-screenshots-hindi.docx"

PAGES = [
    {
        "file": "Screenshot 2026-07-16 171628.png",
        "title": "मुख्य पृष्ठ",
        "desc": "यह वेबसाइट का पहला दृश्य है। यहां मंदिर का नाम, स्थान और मुख्य जानकारी सरल रूप में दिखाई गई है।",
    },
    {
        "file": "Screenshot 2026-07-16 171704.png",
        "title": "सेवाओं का परिचय",
        "desc": "इस भाग में मंदिर की प्रमुख सेवाएं एक साथ दिखाई गई हैं। भक्त यहां से अपनी जरूरत के अनुसार सही विकल्प देख सकते हैं।",
    },
    {
        "file": "Screenshot 2026-07-16 171810.png",
        "title": "भक्ति और संस्कार पृष्ठ",
        "desc": "इस पेज पर कथा, अनुष्ठान और संस्कार से जुड़ी जानकारी दी गई है। सभी विकल्प साफ और पढ़ने में आसान रूप में रखे गए हैं।",
    },
    {
        "file": "Screenshot 2026-07-16 171900.png",
        "title": "नवरात्रि आयोजन पृष्ठ",
        "desc": "यहां नवरात्रि और विशेष पर्वों की जानकारी दिखाई गई है। भक्त संबंधित कार्यक्रमों को आसानी से देख सकते हैं।",
    },
    {
        "file": "Screenshot 2026-07-16 171941.png",
        "title": "मंदिर के बारे में",
        "desc": "इस पेज में मंदिर का परिचय, माहौल और आध्यात्मिक महत्व बताया गया है। साथ ही दृश्य रूप में मंदिर की झलक भी दिखाई गई है।",
    },
    {
        "file": "Screenshot 2026-07-16 172021.png",
        "title": "दान सेवा पृष्ठ",
        "desc": "यह भाग मंदिर को सहयोग देने के लिए बनाया गया है। यहां भक्त दान करने का मुख्य संदेश और संबंधित बटन देख सकते हैं।",
    },
    {
        "file": "Screenshot 2026-07-16 172048.png",
        "title": "संपर्क और नीचे का भाग",
        "desc": "इस पेज में मंदिर का पता, फोन, ईमेल और जरूरी लिंक एक साथ दिए गए हैं। यह जानकारी भक्तों के लिए तुरंत उपयोगी है।",
    },
    {
        "file": "Screenshot 2026-07-16 172124.png",
        "title": "गैलरी पृष्ठ",
        "desc": "यहां मंदिर के चित्र और विशेष झलकियां दिखाई गई हैं। श्रद्धालु इस भाग से मंदिर के माहौल को देख और महसूस कर सकते हैं।",
    },
    {
        "file": "Screenshot 2026-07-16 172213.png",
        "title": "कलश जानकारी फॉर्म",
        "desc": "इस फॉर्म में भक्त अपना नाम, पता और कलश से जुड़ी जानकारी भर सकते हैं। सारी जानकारी आसान तरीके से दर्ज करने की सुविधा दी गई है।",
    },
    {
        "file": "Screenshot 2026-07-16 172242.png",
        "title": "दान विवरण फॉर्म",
        "desc": "इस पेज पर दान करने वाले व्यक्ति की जानकारी और भुगतान से जुड़ी बातें भरी जाती हैं। रूपरेखा सीधी और समझने में सरल रखी गई है।",
    },
    {
        "file": "Screenshot 2026-07-16 172435.png",
        "title": "प्रशासनिक डैशबोर्ड",
        "desc": "यह अंदरूनी प्रबंधन का मुख्य पृष्ठ है। यहां से मंदिर के अलग-अलग कामों की स्थिति और ताजा जानकारी देखी जा सकती है।",
    },
    {
        "file": "Screenshot 2026-07-16 172453.png",
        "title": "कलश पंजीकरण सूची",
        "desc": "इस भाग में तेल, घृत और जवारा पंजीकरण की सूचियां दिखाई जाती हैं। नाम और विवरण खोजने की सुविधा भी यहां उपलब्ध है।",
    },
    {
        "file": "Screenshot 2026-07-16 172518.png",
        "title": "दान प्रबंधन पृष्ठ",
        "desc": "यह पेज मंदिर में आए दान और भुगतान की स्थिति देखने के लिए है। यहां कुल रकम और भुगतान की गिनती साफ रूप में दिखाई गई है।",
    },
    {
        "file": "Screenshot 2026-07-16 172544.png",
        "title": "कार्यक्रम प्रबंधन पृष्ठ",
        "desc": "इस भाग से मंदिर के कार्यक्रम जोड़े और देखे जा सकते हैं। यह पेज आयोजन संबंधी जानकारी को एक जगह रखने में मदद करता है।",
    },
    {
        "file": "Screenshot 2026-07-16 172606.png",
        "title": "गैलरी प्रबंधन पृष्ठ",
        "desc": "यहां मंदिर की तस्वीरों को संभालने और दिखाने की व्यवस्था रखी गई है। सभी चित्रों की स्थिति एक नजर में देखी जा सकती है।",
    },
]


def set_run_font(run, size, bold=False, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Nirmala UI"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Nirmala UI")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Nirmala UI")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Nirmala UI")
    run._element.rPr.rFonts.set(qn("w:cs"), "Nirmala UI")
    if color:
        run.font.color.rgb = color


def add_text(paragraph, text, size, bold=False):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return run


def image_width_for_page(section, image_path):
    usable_width = section.page_width - section.left_margin - section.right_margin
    usable_height = section.page_height - section.top_margin - section.bottom_margin
    reserved_height = Inches(1.7)
    max_width_inches = usable_width / 914400
    max_height_inches = (usable_height - reserved_height) / 914400

    with Image.open(image_path) as img:
        width_px, height_px = img.size

    aspect = width_px / height_px
    width_from_height = max_height_inches * aspect
    return Inches(min(max_width_inches, width_from_height))


def configure_section(section):
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)


def main():
    assert len(PAGES) == 15, "Expected 15 screenshot pages."
    for page in PAGES:
        assert (SCREENSHOT_DIR / page["file"]).exists(), f"Missing file: {page['file']}"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_section(doc.sections[0])

    for index, page in enumerate(PAGES):
        section = doc.sections[-1]

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(4)
        add_text(title, f"{index + 1}. {page['title']}", size=18, bold=True)

        desc = doc.add_paragraph()
        desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        desc.paragraph_format.space_after = Pt(10)
        add_text(desc, page["desc"], size=11)

        image = SCREENSHOT_DIR / page["file"]
        image_paragraph = doc.add_paragraph()
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraph.paragraph_format.space_after = Pt(0)
        image_paragraph.add_run().add_picture(str(image), width=image_width_for_page(section, image))

        if index < len(PAGES) - 1:
            doc.add_page_break()
            configure_section(doc.sections[-1])

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
